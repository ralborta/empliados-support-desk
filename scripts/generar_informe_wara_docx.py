#!/usr/bin/env python3
"""Genera el informe del proyecto Wara/BuilderBot/Odoo en formato Word (.docx).

Mismo contenido que el canvas informe-proyecto-wara-60-dias.
Uso: python3 scripts/generar_informe_wara_docx.py [salida.docx]
"""
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
GREY = RGBColor(0x66, 0x66, 0x66)

timeline = [
    ("17 abr", "Base operativa", "Auth, roles, un hilo abierto por cliente, fusión de duplicados, teléfono WhatsApp canónico.", "3"),
    ("28-29 abr", "Configuración de agente", "Panel para prompt compuesto, respaldo local, fallback manual, limpieza de UI y lectura/escritura del prompt BuilderBot.", "13"),
    ("6 may", "Atilio maestro", "Continuidad ante saludos, estado de caso obligatorio, scripts MCP para sincronizar prompts y corrección de cierre con código de caso.", "6"),
    ("13 may", "Clientes registrados", "Modelo cliente/empresa/patente, importación Excel, edición/borrado, webhook solo para registrados y endpoints BuilderBot/Pulze de contexto.", "20"),
    ("1-4 jun", "Wara base", "Validación por Wara, selección multiempresa, consultas de unidades, odómetro/horómetro, diag, staging/prod y parseo de respuestas reales.", "15"),
    ("5 jun", "Memoria e historial", "Persistencia de mensajes BOT/CUSTOMER, reconstrucción desde tickets, timezone Argentina, prompt de odómetro y endpoint confirmo.", "13"),
    ("7-9 jun", "Mantenimiento, certificados y Odoo", "Mantenimiento operativo, certificados reales, Helpdesk Odoo, confirmaciones flexibles y tickets enriquecidos.", "24"),
    ("10 jun", "Cierre de hardening", "Filtro de newsletters/grupos, Odoo con cliente enriquecido y protección contra patentes de ejemplo en odómetro.", "3"),
]

modules = [
    ("Panel y autenticación", "Listo", "Login por email, roles SUPPORT/ADMIN, sesión, acceso denegado y auth del panel con credenciales Wara/admin."),
    ("Clientes", "Listo", "Esquema empresa/patente, CRUD, importación, borrado en cascada, dashboard y tickets ligados a clientes registrados."),
    ("Tickets internos", "Listo", "Un hilo abierto por cliente, merge de duplicados, tabla, detalle, estado, historial y mensajes entrantes/salientes."),
    ("WhatsApp inbound/outbound", "Listo", "Webhook tolerante a payloads variables, eventos de estado ignorados, mensajes salientes BOT persistidos y no humanos filtrados."),
    ("BuilderBot contexto", "Listo", "Endpoints de cliente registrado, alias de phone/from/key/token, x-api-key, Basic Auth, diagnósticos y hints de secretos."),
    ("BuilderBot Cloud flows", "Activo", "Inicio, Router Wara, selección de empresa, consulta de unidad, odómetro, mantenimiento, certificados, Odoo y BackUp legacy desactivado."),
    ("Wara sesión/multiempresa", "Listo con cuidado", "ObtenerContactosPorNumero, CreateChatBotToken por contacto_id, selección flexible/reset y fallback local ante errores."),
    ("Consulta de unidad", "Listo", "Estado de unidad, último reporte, posición, ignición y alimentación. Ajustado para consulta directa por último reporte."),
    ("Odómetro/horómetro", "Listo", "Registro real en Wara, fecha local, patente con espacios, confirmación flexible, reconstrucción desde historial y filtro de patentes ejemplo."),
    ("Certificados", "Listo", "Endpoint real Certificadocobertura, patente con espacios, confirmación, reconstrucción desde historial y respuesta al cliente."),
    ("Mantenimiento", "Parcial", "Flujo conversacional y trazabilidad local/ticket. API real de mantenimiento sigue dependiente de documentación/endpoint confirmado por Wara."),
    ("Odoo Helpdesk", "Listo", "Diagnóstico, auth JSON-RPC, creación de tickets, partner find/create, equipo Atención al cliente y descripción enriquecida."),
    ("Prompts Atilio", "Activo", "Scope cerrado a soporte Wara, nivel 1 antes de derivar, una pregunta por turno, terminología patente y restricciones anti temas externos."),
    ("Observabilidad/diag", "Activo", "Diag Wara, dumps redactados, flags de sesión, timezone, errores reales de token y validación BuilderBot."),
]

bb_flows = [
    ("Inicio", "Entrada", "EVENTS.WELCOME. Valida el número contra el backend, filtra no humanos, detecta cliente registrado y si debe elegir empresa.", "Ignorar No Cliente, Elegir Empresa, Derivar, Router Wara"),
    ("Mensaje de Voz", "Entrada", "EVENTS.VOICE_NOTE. Transcribe audio y aplica la misma validación de cliente/contexto.", "Elegir Empresa, Derivar, Router Wara"),
    ("Ignorar No Cliente", "Protección", "Silencia por 24h newsletters, broadcasts, grupos o remitentes no humanos para evitar tickets falsos.", "Fin silencioso"),
    ("Derivar", "Derivación", "Mensaje para números no registrados o errores duros de validación. Ya no forwardea mensajes internos al chat.", "Fin / atención manual"),
    ("Elegir Empresa", "Contexto", "Cuando un número tiene varias empresas en Wara, muestra opciones, captura respuesta y guarda la empresa seleccionada.", "Router Wara o Derivar si falla"),
    ("Cambiar Empresa", "Contexto", "Resetea la empresa elegida y vuelve a pedir selección cuando el usuario lo solicita explícitamente.", "Elegir/guardar empresa, luego Router Wara"),
    ("Router Wara", "Router", "Add_intent central. Decide intención: confirmaciones, consulta unidad, odómetro, certificados, mantenimiento, info, Odoo o Atilio.", "Todos los subflujos operativos"),
    ("Flow: Inicio de conversación", "Atilio maestro", "Atención conversacional y soporte nivel 1. Cerrado a temas Wara, diagnostica antes de derivar y pide solo datos pendientes.", "Permanece conversando o deriva según nueva intención"),
    ("Consultar Unidad", "Subflujo conversacional", "Prompt para reunir patente/intención de consulta. Hoy queda como apoyo; la consulta directa va al ejecutor.", "Ejecutar Consulta Unidad"),
    ("Ejecutar Consulta Unidad", "Ejecutor HTTP", "Llama a /api/wara/unidades. Devuelve último reporte, ubicación/posición, ignición y alimentación/voltaje.", "Backend Wara unidades"),
    ("Cambio Odómetro", "Subflujo conversacional", "Reúne patente, odómetro/horómetro y fecha. Resume y pide confirmación natural antes de ejecutar.", "Ejecutar Odómetro"),
    ("Ejecutar Odómetro", "Ejecutor HTTP", "Llama a /api/wara/odometro-horometro con from, rawText y confirmación. Registra el cambio real en Wara.", "Backend Wara odómetro"),
    ("Certificados", "Subflujo operativo", "Pide patente para certificado de cobertura/monitoreo y registra solicitud inicial.", "Ejecutar Certificados o Elegir Empresa"),
    ("Ejecutar Certificados", "Ejecutor HTTP", "Confirma y llama a /api/wara/certificados, conectado al endpoint real Certificadocobertura.", "Backend Wara certificados"),
    ("Gestión Mantenimiento", "Subflujo operativo", "Pide patente, tarea y prioridad. Interpreta/estructura la gestión y deja trazabilidad.", "Ejecutar Gestión Mantenimiento o Elegir Empresa"),
    ("Ejecutar Gestión Mantenimiento", "Ejecutor HTTP", "Confirma la gestión vía /api/wara/mantenimiento-operativo. Hoy es trazabilidad local/ticket por falta de API Wara real confirmada.", "Backend mantenimiento operativo"),
    ("Información Mantenimiento", "FAQ especializada", "ChatPDF con PDF del módulo mantenimiento. Responde dudas conceptuales y rerutea si el usuario quiere ejecutar una acción.", "Odómetro, Consulta, Gestión Mantenimiento o Atilio"),
    ("Reclamo / Asesor (Odoo)", "Ejecutor Odoo", "Captura patente/problema y crea ticket en Odoo Helpdesk, equipo Atención al cliente, con contexto enriquecido.", "Backend /api/odoo/ticket"),
    ("BackUp", "Legacy desactivado", "Flow antiguo con keyword interna backup_legacy_desactivado_wara_internal. Se dejó fuera de interacción normal para evitar dobles respuestas.", "No debería dispararse"),
]

bb_sequence = [
    ("1", "Cliente escribe o manda audio", "Inicio o Mensaje de Voz valida el número con /api/builderbot/customer-registered/:phone/context."),
    ("2", "Filtro inicial", "Si es newsletter/grupo/broadcast va a Ignorar No Cliente. Si no está registrado va a Derivar."),
    ("3", "Contexto de empresa", "Si Wara devuelve varias empresas, Elegir Empresa guarda la opción con /select-company. Si ya hay empresa, sigue."),
    ("4", "Router Wara", "Clasifica intención semántica y contexto: consulta, odómetro, certificado, mantenimiento, Odoo, cambio empresa o Atilio."),
    ("5", "Subflujos conversacionales", "Odómetro, certificados, mantenimiento o Atilio reúnen datos mínimos; consulta unidad ahora puede ir directo al ejecutor."),
    ("6", "Ejecutores HTTP", "Los ejecutores llaman al backend Next.js, que resuelve sesión Wara/Odoo, reconstruye historial si hace falta y devuelve message/summaryText."),
    ("7", "Respuesta y memoria", "El resultado vuelve al WhatsApp y el backend persiste mensajes BOT/CUSTOMER en tickets para continuidad operativa."),
]

functional = [
    ("Wara: validación, sesión y empresa", "Se construyó la capa para resolver el cliente desde WhatsApp, crear sesión Wara por contacto, manejar usuarios con una o varias empresas y permitir reset/cambio de empresa. También se agregaron diagnósticos para distinguir staging/prod, ver respuestas crudas redactadas y entender fallas reales del token."),
    ("Wara: servicios operativos", "Quedaron implementados consulta de unidades, cambio de odómetro/horómetro y certificados reales. En odómetro se ajustaron formato de patente, hora local, confirmación y reconstrucción desde historial. En consulta se expone último reporte, ubicación, ignición y alimentación. En certificados se integró el endpoint real y se corrigió el uso de patente más reciente."),
    ("BuilderBot: ruteo y prompts", "Se separó el Router Wara por intención: consulta, odómetro, mantenimiento, certificados, cambio de empresa, Odoo y Atilio conversacional. Se endureció el prompt maestro para no responder recetas, clima, salud ni temas externos; Atilio ahora opera como soporte nivel 1 y solo deriva cuando corresponde."),
    ("Memoria operativa", "Se agregó persistencia de mensajes salientes del BOT y entrantes del cliente, incluso con payloads variables de BuilderBot. Esto permite reconstruir trámites cuando BB no envía todos los datos en el último mensaje, como confirmaciones de odómetro, mantenimiento o certificados."),
    ("Odoo Helpdesk", "Se creó la integración JSON-RPC con Odoo para tickets de reclamo/escalamiento: diagnóstico, autenticación, partner por teléfono/empresa, creación de ticket, equipo de atención, título y descripción enriquecida con contexto de conversación y datos Wara."),
    ("Protecciones de producción", "Se filtraron newsletters, grupos, broadcasts y remitentes no humanos para evitar tickets falsos. Se removieron respuestas duplicadas por flow legacy, se ajustaron mensajes de faltantes para guiar y no cortar, y se validó/reinició BuilderBot tras cambios críticos."),
]

decisions = [
    ("Derivación a Odoo como último recurso", "Se ajustó la lógica para que Atilio no derive de inmediato. Primero debe entender el problema, pedir datos mínimos, intentar ayuda básica de nivel 1 y solo crear ticket si el caso no se puede resolver automáticamente o si el cliente pide explícitamente un asesor con unidad/patente identificada."),
    ("Scope cerrado a soporte Wara", "Se endurecieron prompts e intents para que Atilio no responda temas ajenos al servicio Wara, como salud, clima, recetas o información general. El fallback debe traer la conversación al soporte Wara."),
    ("Confirmaciones naturales, pero determinísticas", "En servicios críticos se permite lenguaje natural y errores de tipeo, pero el backend reconstruye el trámite desde historial y exige confirmación antes de registrar cambios sensibles como odómetro, mantenimiento o certificados."),
]

pending = [
    ("Mantenimiento API real", "Pendiente externo", "Tenemos flujo y trazabilidad, pero falta endpoint/documentación Wara para registrar mantenimiento 100% real."),
    ("Odoo stage inicial", "Configuración", "Confirmar si se debe forzar etapa inicial específica o dejar default de Helpdesk."),
    ("Pruebas multiempresa reales", "Validación", "Código y flow están preparados, pero faltan pruebas con número real que tenga múltiples empresas."),
    ("Consulta unidad sin patente", "Ajustado recientemente", "Se corrigió para pedir patente/listar unidades sin cerrar la conversación; conviene validar en WhatsApp real."),
    ("Prompts BB Cloud", "Operativo", "Cambios de prompt/intents viven en BuilderBot Cloud; conviene exportar o documentar snapshots si se necesita auditoría formal."),
]


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def main():
    out = sys.argv[1] if len(sys.argv) > 1 else "Informe-Proyecto-Wara-60-dias.docx"
    doc = Document()

    title = doc.add_heading("Informe del proyecto Wara / BuilderBot / Odoo", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(
        "Alcance revisado: últimos 60 días de trabajo del repositorio, con cambios complementarios en BuilderBot Cloud. "
        f"Fuente principal: git log --since=\"60 days ago\". Generado el {date.today().isoformat()}."
    )
    run.font.color.rgb = GREY
    run.font.size = Pt(9)

    doc.add_heading("Indicadores", level=1)
    add_table(
        doc,
        ["Métrica", "Valor"],
        [
            ["Commits últimos 60 días", "96"],
            ["Días con actividad", "14"],
            ["Pico diario de commits", "20 (13 may)"],
            ["Módulos relevados", "14"],
            ["Flows en BuilderBot", "19"],
        ],
    )

    doc.add_heading("Resultado global", level=1)
    doc.add_paragraph(
        "El proyecto pasó de una mesa de ayuda con panel, clientes y tickets a una integración operacional con Wara y "
        "Odoo: validación de clientes por WhatsApp, contexto por empresa, servicios Wara reales, memoria de conversación, "
        "soporte nivel 1 en Atilio y escalamiento a Helpdesk cuando no hay resolución automática."
    )

    doc.add_heading("Línea de tiempo", level=1)
    add_table(doc, ["Fecha", "Bloque", "Qué se hizo", "Commits"], timeline)

    doc.add_heading("Estado por módulo", level=1)
    add_table(doc, ["Módulo", "Estado", "Detalle"], modules)

    doc.add_heading("BuilderBot Cloud: flows activos (19)", level=1)
    doc.add_paragraph(
        "La arquitectura en BuilderBot se organiza en una entrada principal, un router semántico, subflujos de captura, "
        "ejecutores HTTP y flows de protección/contexto."
    )
    add_table(doc, ["Flow", "Tipo", "Función específica", "Conexión real"], bb_flows)

    doc.add_heading("Flujo real entre flows", level=1)
    add_table(doc, ["Paso", "Evento", "Qué pasa realmente"], bb_sequence)

    doc.add_heading("Detalle funcional", level=1)
    for h, body in functional:
        doc.add_heading(h, level=2)
        doc.add_paragraph(body)

    doc.add_heading("Decisiones y ajustes importantes", level=1)
    for h, body in decisions:
        doc.add_heading(h, level=2)
        doc.add_paragraph(body)

    doc.add_heading("Pendientes / riesgos", level=1)
    add_table(doc, ["Tema", "Tipo", "Nota"], pending)

    doc.save(out)
    print(f"OK -> {out}")


if __name__ == "__main__":
    main()
