#!/usr/bin/env python3
"""Genera el documento Word de funcionalidades de Empliados Support Desk."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    out = Path(__file__).resolve().parent.parent / "docs" / "Empliados-Support-Desk-Funcionalidades.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title = doc.add_heading("Empliados Support Desk", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Descripción detallada de todas las funcionalidades de la plataforma"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # --- 1 ---
    add_heading(doc, "1. Propósito y alcance", 1)
    doc.add_paragraph(
        "Empliados Support Desk es una plataforma de mesa de ayuda y soporte para el producto "
        "Empliados. Centraliza las consultas de clientes que llegan principalmente por WhatsApp, "
        "las organiza como tickets numerados y permite al equipo de soporte gestionarlas desde un "
        "panel web, con respuestas automáticas, notificaciones a agentes y asistencia de "
        "inteligencia artificial para resumir conversaciones."
    )

    # --- 2 ---
    add_heading(doc, "2. Tipos de usuarios", 1)
    add_heading(doc, "2.1 Operador del panel", 2)
    doc.add_paragraph(
        "Persona que accede al sistema web con contraseña compartida. Puede ver todos los tickets, "
        "cambiar estados, asignar agentes, escribir respuestas y notas internas, consultar el "
        "dashboard y administrar el listado de agentes."
    )
    add_heading(doc, "2.2 Agentes de soporte", 2)
    doc.add_paragraph(
        "Miembros del equipo registrados con nombre, correo electrónico, teléfono y rol "
        "(Administrador o Soporte). No tienen sesión propia en el panel; reciben notificaciones "
        "por WhatsApp cuando se les asigna un ticket."
    )
    add_heading(doc, "2.3 Clientes", 2)
    doc.add_paragraph(
        "Empresas o contactos identificados por número de teléfono. Envían mensajes por WhatsApp; "
        "el sistema asocia cada conversación a un cliente y a uno o más tickets."
    )

    # --- 3 ---
    add_heading(doc, "3. Acceso y seguridad", 1)
    add_bullets(
        doc,
        [
            "Pantalla de inicio de sesión con contraseña.",
            "Sin inicio de sesión por usuario individual: todos los operadores comparten la misma contraseña de acceso.",
            "Sesión protegida: las rutas del panel requieren estar autenticado.",
            "Cierre de sesión disponible mediante la función correspondiente del sistema (no siempre visible como botón en la interfaz).",
        ],
    )

    # --- 4 ---
    add_heading(doc, "4. Navegación del panel", 1)
    doc.add_paragraph("Menú lateral con las siguientes secciones:")
    add_bullets(
        doc,
        [
            "Inicio: enlace al Dashboard y a la lista de todos los tickets.",
            "Por estado: Abiertos, En progreso, Esperando cliente, Resueltos, Cerrados.",
            "Por prioridad: Urgente, Alta, Normal, Baja.",
            "Gestión: Agentes.",
        ],
    )

    # --- 5 ---
    add_heading(doc, "5. Dashboard (vista general)", 1)
    doc.add_paragraph("Indicadores principales (KPI):")
    add_bullets(
        doc,
        [
            "Total de tickets en el sistema.",
            "Tickets creados en el día actual.",
            "Tickets marcados como resueltos en el día actual.",
            "Tiempo promedio de resolución expresado en horas (calculado sobre tickets resueltos o cerrados).",
        ],
    )
    doc.add_paragraph("Alertas:")
    add_bullets(
        doc,
        [
            "Aviso destacado cuando existen tickets urgentes sin agente asignado, con cantidad y mensaje de atención inmediata.",
        ],
    )
    doc.add_paragraph("Gráficos y listados:")
    add_bullets(
        doc,
        [
            "Distribución de tickets por cada estado (barras de progreso relativas al total).",
            "Distribución de tickets por cada prioridad.",
            "Tendencia de tickets creados en los últimos siete días (un valor por día).",
            "Ranking de los cinco agentes con más tickets activos (no cerrados).",
            "Ranking de las cinco empresas o clientes con más tickets históricos, mostrando nombre y teléfono.",
        ],
    )

    # --- 6 ---
    add_heading(doc, "6. Gestión de tickets — listas", 1)
    doc.add_paragraph(
        "Cada vista de lista muestra hasta cien tickets ordenados por última actividad. "
        "Las columnas incluyen título, cliente, estado, prioridad y datos de asignación según corresponda."
    )
    add_bullets(
        doc,
        [
            "Todos los tickets: vista general con tarjetas resumen (cantidad de abiertos, en progreso, esperando cliente y urgentes).",
            "Abiertos: solo estado Abierto.",
            "En progreso: tickets en atención activa.",
            "Esperando cliente: se envió respuesta y se espera respuesta del cliente.",
            "Resueltos: incluye tres contadores — total resueltos, resueltos por IA y resueltos por humano.",
            "Cerrados: tickets archivados.",
            "Por prioridad (Urgente, Alta, Normal, Baja): filtro exclusivo por nivel de prioridad.",
        ],
    )

    # --- 7 ---
    add_heading(doc, "7. Detalle de un ticket", 1)
    add_heading(doc, "7.1 Información mostrada", 2)
    add_bullets(
        doc,
        [
            "Código único del ticket (identificador legible para cliente y equipo).",
            "Estado actual, prioridad y categoría (Soporte técnico, Facturación, Ventas, Otro).",
            "Canal (por defecto WhatsApp; el modelo admite Email y Web).",
            "Título del caso.",
            "Datos del cliente: empresa, nombre de contacto, teléfono.",
            "Agente asignado (si existe).",
        ],
    )
    add_heading(doc, "7.2 Conversación", 2)
    add_bullets(
        doc,
        [
            "Historial cronológico de mensajes.",
            "Cada mensaje indica si proviene del cliente, del bot o de un agente humano.",
            "Mensajes de solo lectura para el equipo cuando corresponde a notas internas frente a respuestas al cliente.",
            "Visualización de archivos adjuntos (imágenes, videos, documentos) vinculados a cada mensaje.",
        ],
    )
    add_heading(doc, "7.3 Resumen de conversación (IA)", 2)
    add_bullets(
        doc,
        [
            "Bloque de resumen generado por inteligencia artificial (texto breve del problema y contexto).",
            "Botón para generar o regenerar el resumen manualmente.",
            "Indicación de que el resumen puede actualizarse automáticamente al añadir mensajes.",
        ],
    )
    add_heading(doc, "7.4 Asignación de agente", 2)
    add_bullets(
        doc,
        [
            "Selector para asignar o cambiar el agente responsable del ticket.",
            "Al asignar, el agente recibe un mensaje por WhatsApp con datos del ticket, prioridad, estado, resumen y enlace al panel.",
        ],
    )
    add_heading(doc, "7.5 Cambio de estado", 2)
    add_bullets(
        doc,
        [
            "Esperando cliente: marca que se espera respuesta del cliente (el botón se deshabilita si ya está en ese estado).",
            "Marcar como resuelto: indica que el caso quedó resuelto.",
            "Cerrar: archiva el ticket (el botón se deshabilita si ya está cerrado).",
        ],
    )
    add_heading(doc, "7.6 Composición de mensajes", 2)
    add_bullets(
        doc,
        [
            "Campo de texto para escribir contenido.",
            "Modo Respuesta al cliente: el mensaje se envía al número de WhatsApp del cliente y queda registrado en la conversación.",
            "Modo Nota interna: visible solo para el equipo; el cliente no la recibe.",
            "Botones Guardar nota o Responder según el modo seleccionado.",
        ],
    )
    doc.add_paragraph(
        "Tras enviar una respuesta al cliente, el sistema puede actualizar automáticamente el "
        "resumen por IA con el nuevo contexto de la conversación."
    )

    # --- 8 ---
    add_heading(doc, "8. Canal WhatsApp — recepción de mensajes", 1)
    add_bullets(
        doc,
        [
            "Los mensajes entrantes se procesan cuando el proveedor de chatbot notifica un nuevo mensaje.",
            "Solo se procesan eventos de mensaje entrante; otros eventos se ignoran de forma controlada.",
            "Identificación del cliente por número de teléfono; creación o actualización del registro de cliente.",
            "Evita duplicados: si el mismo mensaje se reenvía, no se registra dos veces.",
            "Acepta mensaje solo con archivos adjuntos (sin texto), usando un texto placeholder cuando hace falta.",
        ],
    )
    add_heading(doc, "8.1 Formato de mensaje recomendado", 2)
    add_bullets(
        doc,
        [
            "Si el mensaje tiene al menos tres líneas no vacías: línea 1 = nombre de empresa, línea 2 = nombre y rol del contacto, líneas siguientes = consulta o problema.",
            "Si no sigue ese formato, se usa el texto completo como consulta.",
        ],
    )
    add_heading(doc, "8.2 Creación o reutilización de ticket", 2)
    add_bullets(
        doc,
        [
            "Se busca un ticket abierto reciente del mismo cliente (ventana de 48 horas desde el último mensaje).",
            "Si existe, se agrega el mensaje a ese ticket.",
            "Si no, se crea un ticket nuevo con código único, título derivado del inicio del mensaje y estado inicial abierto.",
        ],
    )
    add_heading(doc, "8.3 Archivos adjuntos", 2)
    add_bullets(
        doc,
        [
            "Los archivos que envía el cliente se almacenan de forma segura y se vinculan al mensaje con URL, tipo y nombre.",
        ],
    )
    add_heading(doc, "8.4 Clasificación automática", 2)
    add_bullets(
        doc,
        [
            "Prioridad y categoría se infieren por palabras clave en el texto (por ejemplo urgencia, producción, facturación, etc.).",
        ],
    )
    add_heading(doc, "8.5 Escalación automática", 2)
    add_bullets(
        doc,
        [
            "Puede activarse por prioridad urgente, por palabras sensibles (amenazas legales, fraude, cliente molesto, etc.) o cuando el ticket ya acumula tres o más mensajes previos.",
            "En caso de escalación, el mensaje automático al cliente indica que el caso fue escalado al equipo.",
            "Si no hay escalación, el mensaje automático confirma recepción y muestra el código del ticket.",
        ],
    )

    # --- 9 ---
    add_heading(doc, "9. Gestión de agentes", 1)
    add_bullets(
        doc,
        [
            "Listado de todos los agentes con nombre, correo, teléfono, rol y cantidad de tickets asignados.",
            "Alta de agente: formulario con nombre, email, teléfono (formato tipo número internacional) y rol Administrador o Soporte.",
            "Eliminación de agente: permitida solo si no tiene tickets asignados; en caso contrario el sistema impide borrar y muestra la razón.",
        ],
    )

    # --- 10 ---
    add_heading(doc, "10. Modelo de datos de negocio (conceptos)", 1)
    add_bullets(
        doc,
        [
            "Cliente: teléfono único, nombre opcional, relación con todos sus tickets.",
            "Ticket: código, cliente, contacto, título, estado, prioridad, categoría, canal, agente asignado, fechas de creación y última actividad.",
            "Campos de IA en el ticket: resumen de conversación, texto de resolución o conclusión, indicador de si fue resuelto por IA.",
            "Mensaje: dirección (entrante, saliente, nota interna), origen (cliente, bot, humano), texto, adjuntos, identificador externo para integración.",
            "Eventos de ticket: registro de cambios de estado, asignaciones, etiquetas, prioridad, respuestas automáticas y escalaciones (para auditoría e integraciones).",
            "Etiquetas (tags): el modelo permite etiquetar tickets; la interfaz actual no expone gestión de etiquetas.",
        ],
    )

    # --- 11 ---
    add_heading(doc, "11. Procesos automáticos internos (sin pantalla dedicada)", 1)
    doc.add_paragraph(
        "La plataforma incluye capacidades que pueden invocarse desde integraciones o procesos "
        "automatizados, no desde botones visibles en el panel para el operador habitual:"
    )
    add_bullets(
        doc,
        [
            "Escalación formal del ticket: genera resumen y conclusión por IA, actualiza el ticket y puede eliminar mensajes temporales según la lógica del proceso.",
            "Cierre por IA: similar al anterior, orientado a cerrar casos tratados automáticamente, dejando resumen y conclusión registrados.",
        ],
    )

    # --- 12 ---
    add_heading(doc, "12. Limitaciones y matices", 1)
    add_bullets(
        doc,
        [
            "Los roles de agente (Administrador vs Soporte) no cambian las pantallas disponibles en el panel para el operador.",
            "Las etiquetas existen en el modelo de datos pero no hay pantalla para gestionarlas en la versión actual.",
            "La lista general de tickets está limitada a los últimos cien registros por vista.",
            "No se encontraron comentarios TODO pendientes en el código fuente de referencia; la lista anterior refleja el comportamiento implementado.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Resumen", 1)
    doc.add_paragraph(
        "En conjunto, Empliados Support Desk permite recibir consultas por WhatsApp, convertirlas en "
        "tickets trazables, priorizarlas y clasificarlas automáticamente, responder con mensajes "
        "automáticos y manuales, asignar trabajo a agentes con notificación móvil, documentar la "
        "conversación con resúmenes asistidos por IA y supervisar la operación mediante dashboard "
        "y listados filtrados por estado y prioridad."
    )

    doc.save(out)
    print(f"Documento guardado en: {out}")


if __name__ == "__main__":
    main()
