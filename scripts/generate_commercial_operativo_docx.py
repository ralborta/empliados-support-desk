#!/usr/bin/env python3
"""Documento comercial y operativo — Empliados Support Desk (Word)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    out = (
        Path(__file__).resolve().parent.parent
        / "docs"
        / "Empliados-Support-Desk-Descripcion-Comercial-y-Operativa.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    t = doc.add_heading("Empliados Support Desk", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(
        "Ficha de producto — descripción comercial y operativa"
    )
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = doc.add_paragraph(
        "Documento orientado a presentación comercial, planificación operativa y replicación del modelo en otros proyectos."
    )
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.runs[0].italic = True
    doc.add_paragraph()

    # --- Resumen ejecutivo ---
    add_heading(doc, "1. Resumen ejecutivo", 1)
    doc.add_paragraph(
        "Empliados Support Desk es una solución de mesa de ayuda orientada a empresas que "
        "atienden a sus clientes por WhatsApp. Convierte cada conversación en un caso "
        "gestionable (ticket), con priorización, seguimiento por estados, asignación a "
        "personas del equipo, respuestas automáticas al cliente y apoyo de inteligencia "
        "artificial para resumir conversaciones. El equipo de soporte trabaja desde un panel "
        "web unificado; los clientes no necesitan instalar nada: siguen usando WhatsApp como "
        "hoy."
    )

    # --- Propuesta de valor ---
    add_heading(doc, "2. Propuesta de valor comercial", 1)
    add_bullets(
        doc,
        [
            "Un solo lugar para ver todas las consultas, sin depender de chats dispersos en teléfonos personales.",
            "Trazabilidad: cada caso tiene código, historial y estado visible para el equipo.",
            "Respuesta inmediata al cliente con acuse de recibo y número de ticket, mejorando la percepción de servicio.",
            "Priorización automática que ayuda a atender antes lo urgente o sensible.",
            "Visibilidad gerencial: volumen de trabajo, tiempos y carga por agente y por cliente.",
            "Escalación asistida cuando el volumen de mensajes o el tipo de consulta lo requieren.",
        ],
    )

    # --- Problema / oportunidad ---
    add_heading(doc, "3. Problema que aborda y oportunidad", 1)
    doc.add_paragraph(
        "Muchas organizaciones reciben soporte por WhatsApp de forma informal: se pierde "
        "contexto, no hay métricas claras y es difícil repartir trabajo o demostrar niveles de "
        "servicio. Esta plataforma formaliza ese canal sin cambiar el hábito del cliente "
        "(sigue en WhatsApp) mientras el negocio gana estructura, medición y colaboración."
    )

    # --- Para quién es ---
    add_heading(doc, "4. Público objetivo y casos de uso", 1)
    add_heading(doc, "4.1 Perfiles", 2)
    add_bullets(
        doc,
        [
            "Equipos de soporte técnico, facturación o ventas que ya usan WhatsApp con clientes B2B o B2C.",
            "Empresas de software o servicios (como Empliados) con volumen medio de consultas y necesidad de orden y seguimiento.",
            "Operaciones que requieren dejar registro de conversaciones para calidad o auditoría.",
        ],
    )
    add_heading(doc, "4.2 Casos de uso típicos", 2)
    add_bullets(
        doc,
        [
            "Cliente reporta un incidente o duda; recibe confirmación con número de ticket y el equipo lo gestiona hasta el cierre.",
            "Consulta sensible o repetida: el sistema puede escalar el caso para que un humano lo priorice.",
            "Supervisor revisa dashboard para ver carga del día, cuellos de botella y clientes más demandantes.",
        ],
    )

    # --- Características detalladas ---
    add_heading(doc, "5. Características y funcionalidades del producto", 1)
    doc.add_paragraph(
        "A continuación se listan las capacidades actuales, agrupadas por ámbito. Sirven como "
        "inventario para comercialización, contratos de alcance o clonación del modelo en otro proyecto."
    )

    add_heading(doc, "5.1 Canal de entrada: WhatsApp", 2)
    add_bullets(
        doc,
        [
            "Recepción de mensajes de texto desde el canal conectado al negocio.",
            "Recepción de archivos adjuntos (imágenes, video, documentos) asociados al mensaje.",
            "Identificación automática del cliente por número de teléfono.",
            "Formato de mensaje estructurado opcional: empresa, contacto y descripción en líneas separadas para mejor datos maestros.",
            "Si el cliente no sigue el formato, el mensaje completo se trata como consulta.",
            "Reapertura inteligente: mensajes recientes del mismo cliente en ventana de 48 horas se agrupan en el mismo ticket en curso.",
            "Generación de código único de ticket legible para cliente y equipo.",
            "Respuestas automáticas al cliente: confirmación de recepción con número de ticket y, si aplica, aviso de escalación al equipo.",
        ],
    )

    add_heading(doc, "5.2 Clasificación y priorización automática", 2)
    add_bullets(
        doc,
        [
            "Inferencia de prioridad (baja, normal, alta, urgente) según palabras clave en el mensaje.",
            "Categorización automática (p. ej. soporte técnico, facturación, ventas, otro).",
            "Detección de situaciones que ameritan escalación: prioridad urgente, vocabulario de riesgo o conflicto, o ticket con varios mensajes acumulados.",
        ],
    )

    add_heading(doc, "5.3 Panel web para el equipo de soporte", 2)
    add_bullets(
        doc,
            [
            "Acceso protegido por contraseña de equipo (acceso operativo centralizado).",
            "Vista de listado de tickets con orden por última actividad.",
            "Filtros por estado: abiertos, en progreso, esperando respuesta del cliente, resueltos, cerrados.",
            "Filtros por prioridad: urgente, alta, normal, baja.",
            "Vista consolidada de todos los tickets con mini-resumen de conteos clave.",
            "Detalle de ticket con ficha de cliente (empresa, contacto, teléfono), estado, prioridad, categoría y canal.",
            "Hilo de conversación con distinción visual entre mensajes del cliente, del sistema/bot y del agente.",
            "Notas internas: texto visible solo para el equipo, sin envío al cliente.",
            "Respuestas al cliente desde el panel que se entregan por el mismo canal WhatsApp.",
            "Cambio de estado operativo: esperando cliente, resuelto, cerrado.",
            "Asignación y reasignación de tickets a miembros del equipo registrados como agentes.",
        ],
    )

    add_heading(doc, "5.4 Inteligencia artificial (asistencia)", 2)
    add_bullets(
        doc,
        [
            "Resumen automático de la conversación en pocas líneas (problema, datos clave, urgencia).",
            "Generación o actualización manual del resumen bajo demanda.",
            "Actualización del resumen cuando entran nuevos mensajes relevantes (según la lógica del producto).",
            "Capacidad de procesos de cierre o escalación asistidos que consolidan resumen y conclusión (pensados para integraciones o automatizaciones).",
            "En la vista de tickets resueltos, distinción entre resolución atribuible a flujo automatizado/IA frente a intervención humana (indicadores operativos).",
        ],
    )

    add_heading(doc, "5.5 Gestión de agentes y notificaciones", 2)
    add_bullets(
        doc,
        [
            "Alta de agentes con nombre, correo, teléfono y rol (administrador o soporte).",
            "Listado con carga de trabajo (tickets asignados no cerrados).",
            "Eliminación de agentes cuando no tienen casos asignados.",
            "Notificación por WhatsApp al agente cuando se le asigna un ticket, incluyendo datos del caso y enlace al panel.",
        ],
    )

    add_heading(doc, "5.6 Indicadores y supervisión (dashboard)", 2)
    add_bullets(
        doc,
        [
            "Total histórico de tickets.",
            "Tickets creados en el día y tickets resueltos en el día.",
            "Tiempo promedio de resolución en horas (referencia operativa).",
            "Alerta destacada de tickets urgentes sin asignar.",
            "Distribución por estado y por prioridad.",
            "Tendencia de tickets creados en los últimos siete días.",
            "Ranking de agentes con mayor carga activa.",
            "Ranking de empresas o clientes con más tickets (volumen y contacto).",
        ],
    )

    add_heading(doc, "5.7 Modelo de negocio en datos (trazabilidad)", 2)
    add_bullets(
        doc,
        [
            "Registro de cliente con teléfono como identificador principal.",
            "Historial de mensajes y metadatos de integración para evitar duplicados.",
            "Eventos de ciclo de vida del ticket (cambios de estado, asignación, prioridad, escalación, respuestas automáticas) para análisis o integraciones futuras.",
            "Soporte en modelo de datos para etiquetado de tickets (extensible; la interfaz actual puede evolucionar).",
        ],
    )

    # --- Operativo día a día ---
    add_heading(doc, "6. Descripción operativa (día a día)", 1)
    add_heading(doc, "6.1 Flujo del cliente", 2)
    add_bullets(
        doc,
        [
            "El cliente escribe por WhatsApp.",
            "Recibe confirmación con su número de ticket.",
            "Si el caso se escala, recibe un mensaje que lo indica.",
            "Recibe respuestas humanas cuando el equipo contesta desde el panel.",
        ],
    )
    add_heading(doc, "6.2 Flujo del agente o responsable", 2)
    add_bullets(
        doc,
        [
            "Ingresa al panel y prioriza por listas (urgentes, abiertos, etc.) o por dashboard.",
            "Abre el ticket, lee resumen IA y conversación.",
            "Asigna o se auto-asigna según política del equipo.",
            "Responde al cliente o deja nota interna.",
            "Marca estados conforme avanza el caso hasta resolverlo y cerrarlo.",
        ],
    )
    add_heading(doc, "6.3 Flujo del supervisor", 2)
    add_bullets(
        doc,
        [
            "Revisa KPIs y alertas de urgentes sin asignar.",
            "Identifica clientes de alto volumen y reparte carga según ranking de agentes.",
            "Utiliza tiempos y tendencia semanal para planificar turnos o refuerzos.",
        ],
    )

    # --- Comercial: mensajes clave ---
    add_heading(doc, "7. Mensajes clave para uso comercial", 1)
    add_bullets(
        doc,
        [
            "“Su cliente sigue en WhatsApp; su equipo trabaja con tickets y métricas.”",
            "“Cada consulta tiene número de caso y historial.”",
            "“Prioridad y categoría ayudan a no perder lo urgente.”",
            "“Dashboard listo para reuniones de servicio y mejora continua.”",
        ],
    )

    # --- Clonar / otro proyecto ---
    add_heading(doc, "8. Notas para replicar el modelo en otro proyecto", 1)
    doc.add_paragraph(
        "Si se clona o adapta este producto a otra marca o vertical, lo funcional que debe "
        "replicarse o redefinirse suele ser:"
    )
    add_bullets(
        doc,
        [
            "Nombre de producto, tono de mensajes automáticos al cliente y política de ventana de reutilización de ticket (p. ej. 48 horas).",
            "Reglas de palabras clave para prioridad, categoría y escalación, alineadas al negocio.",
            "Roles y nombres de equipo (agentes) y si las notificaciones siguen siendo por WhatsApp u otro canal.",
            "KPIs del dashboard que el cliente final valore (SLA, satisfacción, etc.) como evolución natural.",
            "Política de acceso al panel (contraseña única vs. usuarios individuales) según cumplimiento y escala.",
        ],
    )
    doc.add_paragraph(
        "Este apartado no prescribe tecnología: describe qué hay que decidir al portar la "
        "experiencia comercial y operativa a otro contexto."
    )

    # --- Cierre ---
    add_heading(doc, "9. Cierre", 1)
    doc.add_paragraph(
        "Empliados Support Desk ofrece una propuesta coherente: canal conversacional popular "
        "(WhatsApp), operación ordenada (tickets y estados), eficiencia (automatización e IA "
        "de resumen) y control (dashboard y asignaciones). El documento resume las "
        "características actuales para apoyo comercial, operativo y de planificación de "
        "producto en nuevos proyectos."
    )

    doc.save(out)
    print(f"Guardado: {out}")


if __name__ == "__main__":
    main()
