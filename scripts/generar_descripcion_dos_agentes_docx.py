#!/usr/bin/env python3
"""Genera documento Word: funcionalidades del servicio Atilio/Wara y división en dos agentes.

Uso: python3 scripts/generar_descripcion_dos_agentes_docx.py [salida.docx]
"""
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
GREY = RGBColor(0x66, 0x66, 0x66)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main():
    root = Path(__file__).resolve().parent.parent
    default_out = root / "docs" / "Atilio-Wara-Descripcion-Servicio-Dos-Agentes.docx"
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else default_out
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = doc.add_heading("Atilio — Mesa de Ayuda Wara", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Descripción de funcionalidades del servicio y propuesta de división en dos agentes"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    meta = doc.add_paragraph(f"Versión: {date.today().strftime('%d/%m/%Y')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in meta.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
    doc.add_paragraph()

    add_heading(doc, "1. Resumen ejecutivo", 1)
    doc.add_paragraph(
        "El servicio Atilio integra WhatsApp (BuilderBot Cloud), un backend Next.js en Vercel "
        "(wara.nivel41.com), la API operativa de Wara, tickets locales en PostgreSQL (Railway) "
        "y escalamiento a Odoo Helpdesk. El cliente interactúa por WhatsApp; Mesa de Ayuda "
        "opera desde el panel interno y Odoo."
    )
    doc.add_paragraph(
        "Este documento resume todas las funcionalidades actuales y propone dividir el sistema "
        "en dos roles lógicos — Agente de Atención (interacción con el cliente) y Agente de "
        "Coordinación Interna (ejecución técnica, reglas y tickets) — sin necesidad de "
        "duplicar bots ni infraestructura WhatsApp."
    )

    add_heading(doc, "2. Arquitectura general", 1)
    add_bullets(
        doc,
        [
            "Canal cliente: WhatsApp vía BuilderBot Cloud (BBC).",
            "Aplicación: Next.js en Vercel (APIs Wara, Odoo, contexto, webhooks).",
            "Base de datos: PostgreSQL en Railway (tickets, clientes, mensajes, sesiones).",
            "Integraciones: API Wara (unidades, certificados, odómetro), Odoo Helpdesk.",
            "Panel web: gestión de tickets, agentes, clientes, prompts y PDFs de guía.",
        ],
    )

    add_heading(doc, "3. Funcionalidades del servicio por capas", 1)

    add_heading(doc, "3.1 Capa de entrada y contexto", 2)
    add_table(
        doc,
        ["Función", "Descripción"],
        [
            ("Identificación", "Valida teléfono en Wara, nombre del contacto y empresa(s) asociadas."),
            ("Multiempresa", "Menú de opciones, elegir/cambiar/reiniciar empresa Wara activa."),
            ("Saludo y continuidad", "Atilio se identifica; retoma trámite sin exponer tickets internos al azar."),
            ("Router semántico", "Clasifica intención del mensaje (~20 reglas en BBC Router Wara)."),
            ("Memoria conversacional", "Hilo persistido en base de datos + historial en prompts BBC."),
        ],
    )

    add_heading(doc, "3.2 Trámites operativos (Wara en vivo)", 2)
    doc.add_paragraph(
        "Requieren sesión Wara válida y consultan datos reales de la plataforma. "
        "El backend resuelve la unidad por patente, marca o nombre (IA + reglas sobre la flota)."
    )
    add_table(
        doc,
        ["Trámite", "Qué pide el cliente", "API backend", "Resultado"],
        [
            (
                "Consulta unidad / GPS",
                "Reporte, offline, ubicación, ignición, listado de flota",
                "/api/wara/unidades",
                "Texto al cliente; ticket automático si aplica Caso 1, 2 o 3",
            ),
            (
                "Certificado de cobertura",
                "Constancia de monitoreo por unidad",
                "/api/wara/certificados",
                "Enlace/PDF o mensaje de error; ticket si falla validación",
            ),
            (
                "Odómetro / horómetro",
                "Registrar kilómetros u horas de motor",
                "/api/wara/odometro-horometro",
                "Registro en Wara tras confirmación (CONFIRMO)",
            ),
            (
                "Mantenimiento operativo",
                "Preventivo, correctivo, tarea con patente y detalle",
                "/api/wara/mantenimiento-operativo",
                "Trazabilidad local y ticket Odoo",
            ),
        ],
    )
    doc.add_paragraph(
        "Monitoreo GPS (consulta unidad): el backend aplica el flujograma Mesa de Ayuda Wara "
        "(falta de reporte, pérdida de señal, falla de ignición, unidad detenida sin ticket) "
        "y enriquece tickets Odoo con telemetría (reporte, posición, ignición, voltaje, coordenadas)."
    )

    add_heading(doc, "3.3 Guías de plataforma (informativas)", 2)
    doc.add_paragraph(
        "No consultan API operativa ni generan tickets por dudas de navegación. "
        "Utilizan ChatPDF en BuilderBot con documentación cargada desde el panel."
    )
    add_table(
        doc,
        ["Módulo", "Temas", "Flow BBC"],
        [
            ("Información Opciones", "Agenda, perfiles, notificaciones, alertas, permisos", "Info Opciones"),
            ("Información Unidades", "Panel de flota, grupos, MIS ATAJOS, puntos de color", "Info Unidades"),
            ("Información Mantenimiento", "Uso del módulo mantenimiento (sin abrir caso real)", "Info Mantenimiento"),
        ],
    )

    add_heading(doc, "3.4 Escalamiento a atención humana", 2)
    add_bullets(
        doc,
        [
            "Flow Reclamo / Asesor → creación de ticket en Odoo Helpdesk.",
            "Tickets automáticos desde consulta GPS (fallas técnicas detectadas).",
            "Tickets desde certificados o mantenimiento cuando la validación o el negocio lo requieren.",
            "Agentes humanos responden desde el panel; mensajes salen por API BuilderBot al WhatsApp del cliente.",
        ],
    )

    add_heading(doc, "3.5 Coordinación interna (panel y Odoo)", 2)
    add_bullets(
        doc,
        [
            "Webhook WhatsApp: persiste mensajes entrantes/salientes, audio (transcripción), adjuntos.",
            "Tickets locales: código TCK, estados, un hilo abierto por cliente, threading.",
            "Sincronización Odoo: partner, equipo Atención al cliente, descripción enriquecida.",
            "IA interna: resumen de conversación, sugerencia de cierre, escalación.",
            "Configuración: prompts por módulo, PDFs de guía, clientes, importación Excel, estadísticas.",
        ],
    )

    add_heading(doc, "4. Flows BuilderBot Cloud (referencia)", 1)
    add_table(
        doc,
        ["Flow", "Rol"],
        [
            ("Inicio / Mensaje de voz", "Validación de cliente y contexto de empresa"),
            ("Router Wara", "Clasificación de intención hacia subflujos"),
            ("Elegir / Cambiar empresa", "Multiempresa Wara"),
            ("Consultar Unidad + Ejecutar Consulta", "Captura conversacional + HTTP unidades"),
            ("Odómetro + Ejecutar Odómetro", "Captura + registro Wara"),
            ("Certificados + Ejecutar Certificados", "Captura + emisión certificado"),
            ("Gestión Mantenimiento + Ejecutar", "Captura + trazabilidad/ticket"),
            ("Info Opciones / Unidades / Mantenimiento", "Guías ChatPDF"),
            ("Reclamo / Asesor Odoo", "Escalamiento humano"),
            ("Derivar / Ignorar", "No registrados o remitentes no humanos"),
        ],
    )

    add_heading(doc, "5. Propuesta: dos agentes lógicos", 1)
    doc.add_paragraph(
        "No se recomienda crear dos bots WhatsApp separados. La división propuesta es por "
        "responsabilidad: quién habla con el cliente versus quién ejecuta reglas y coordina "
        "internamente. En la implementación actual esto ya se acerca al modelo "
        "«BBC delgado + backend grueso»."
    )

    add_heading(doc, "5.1 Agente 1 — Atención / Interacción (Atilio Front)", 2)
    doc.add_paragraph("Canal: WhatsApp (BuilderBot Cloud). Audiencia: cliente final.")
    doc.add_paragraph("Misión:")
    add_bullets(
        doc,
        [
            "Saludar, identificarse como Atilio de Mesa de Ayuda Wara.",
            "Gestionar selección y cambio de empresa cuando corresponde.",
            "Entender la intención del cliente y enrutar al trámite correcto.",
            "Conducir guías informativas (Opciones, Unidades, Mantenimiento).",
            "Capturar datos mínimos: patente, confirmación natural («¿avanzo?», «dale», «sí»).",
            "Invocar el backend y mostrar al cliente el texto devuelto (summaryText / message).",
        ],
    )
    doc.add_paragraph("Límites — NO debe:")
    add_bullets(
        doc,
        [
            "Inventar reporte, posición, ignición, voltaje ni ubicación.",
            "Decidir Caso 1 / 2 / 3 de GPS ni prioridad de tickets.",
            "Abrir o modificar tickets Odoo directamente.",
            "Validar flota ni aplicar reglas de timestamps de telemetría.",
        ],
    )

    add_heading(doc, "5.2 Agente 2 — Coordinación Interna (Atilio Ops)", 2)
    doc.add_paragraph(
        "Canal: APIs REST (Next.js), panel web, Odoo. Audiencia: sistema y Mesa de Ayuda."
    )
    doc.add_paragraph("Misión:")
    add_bullets(
        doc,
        [
            "Resolver unidad por patente, marca o nombre parcial (IA + flota Wara).",
            "Ejecutar consultas y trámites contra API Wara.",
            "Aplicar reglas GPS (flujograma + cruces de timestamps, unidad detenida, falla de ignición).",
            "Crear y enriquecer tickets (local TCK + Odoo) con telemetría y motivo técnico.",
            "Persistir conversación, webhooks, resúmenes IA y escalación a agentes humanos.",
            "Redactar respuesta final al cliente (plantilla o IA anclada a hechos verificados).",
        ],
    )
    doc.add_paragraph("Límites — NO debe:")
    add_bullets(
        doc,
        [
            "Mantener small talk ni saludos repetitivos con el cliente.",
            "Sustituir guías ChatPDF de módulos de plataforma.",
            "Exponer lógica de router de 20 intenciones en prompts largos de BBC.",
        ],
    )

    add_heading(doc, "6. Reparto por trámite", 1)
    add_table(
        doc,
        ["Trámite", "Agente 1 (Atención)", "Agente 2 (Coordinación)"],
        [
            ("Saludo / empresa", "Sí", "—"),
            ("Consulta GPS / unidades", "Captura y muestra respuesta", "Reglas, Wara, tickets"),
            ("Certificado", "Captura y muestra respuesta", "Wara, validación flota"),
            ("Odómetro / horómetro", "Captura y CONFIRMO", "Registro en Wara"),
            ("Mantenimiento operativo", "Captura y CONFIRMO", "Ticket Odoo, trazabilidad"),
            ("Guías Opciones / Unidades / Mant.", "Sí (ChatPDF)", "—"),
            ("Asesor / reclamo", "Deriva", "Ticket Odoo + panel"),
            ("Respuesta agente humano", "—", "Panel → WhatsApp"),
        ],
    )

    add_heading(doc, "7. Flujo de interacción recomendado", 1)
    doc.add_paragraph(
        "1. Cliente escribe por WhatsApp.\n"
        "2. Agente 1 (BBC): valida contexto, enruta, captura datos si faltan.\n"
        "3. Agente 2 (backend): consulta Wara, aplica reglas, decide ticket sí/no.\n"
        "4. Agente 2 devuelve summaryText con hechos verificados.\n"
        "5. Agente 1 muestra el mensaje al cliente.\n"
        "6. Si hay ticket, Agente 2 lo registra en Odoo; Mesa de Ayuda actúa desde el panel."
    )

    add_heading(doc, "8. Evolución futura (opcional)", 1)
    add_bullets(
        doc,
        [
            "Migración a BuilderBot Open Source con bot delgado (mismo reparto de roles).",
            "Router centralizado en backend en lugar de ~20 reglas en BBC.",
            "Dos prompts IA explícitos: Front (solo conversación) y Ops (solo ejecución).",
            "No duplicar bots WhatsApp: un canal, dos capas lógicas.",
        ],
    )

    add_heading(doc, "9. Conclusión", 1)
    doc.add_paragraph(
        "El servicio Atilio/Wara cubre cinco bloques funcionales: contexto del cliente, "
        "cuatro trámites operativos con API Wara, tres guías informativas, escalamiento humano "
        "y coordinación interna (panel + Odoo). La división natural no es por tipo de trámite "
        "sino por audiencia: cliente versus Mesa de Ayuda."
    )
    doc.add_paragraph(
        "La arquitectura actual — BuilderBot para conversación y enrutamiento liviano, "
        "backend Next.js para inteligencia operativa, reglas GPS y tickets — ya implementa "
        "esta separación. Reforzar este modelo es preferible a mantener prompts extensos en "
        "BBC o a desplegar dos bots en paralelo."
    )

    doc.save(out)
    print(f"Generado: {out}")


if __name__ == "__main__":
    main()
